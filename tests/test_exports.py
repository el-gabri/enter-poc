"""Tests for PDF/DOCX conversion from the canonical Markdown."""

import io

from app.reporting.convert import parse_blocks, render_docx, render_pdf

SAMPLE_MD = """# Titulo

## Secao

Paragrafo com **negrito** no meio.

- item um
- item dois com **enfase**

> aviso importante

---
"""


def test_parse_blocks_structure() -> None:
    kinds = [b.kind for b in parse_blocks(SAMPLE_MD)]
    assert kinds == ["h1", "h2", "paragraph", "bullet", "bullet", "quote", "rule"]


def test_docx_roundtrip() -> None:
    import docx

    payload = render_docx(SAMPLE_MD)
    document = docx.Document(io.BytesIO(payload))
    texts = [p.text for p in document.paragraphs if p.text]
    assert "Titulo" in texts
    assert "Paragrafo com negrito no meio." in texts
    # bold run preserved
    bold_runs = [
        run.text
        for p in document.paragraphs
        for run in p.runs
        if run.bold
    ]
    assert "negrito" in bold_runs


def test_pdf_is_valid() -> None:
    payload = render_pdf(SAMPLE_MD)
    assert payload.startswith(b"%PDF")
    assert len(payload) > 1000
